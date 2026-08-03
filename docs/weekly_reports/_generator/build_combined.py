"""Build one Word file: 44 weekly reports, one page per week.

Run from repo root or this folder:
  python3 docs/weekly_reports/_generator/build_combined.py

Output:
  docs/weekly_reports/ALL_WEEKS_COMBINED.docx

Then in Word: File -> Save As -> PDF (44 pages expected).

Progress/plan text is auto-shortened so each week fits one page.
"""
import re
import sys
from copy import deepcopy
from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Inches, Pt
from docx.table import Table

HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(HERE))
from plan_data import WEEKS

TEMPLATE = Path("/u/madhus/Week 1 (002).docx")
SIGNATURE_CANDIDATES = [
    Path(
        "/u/madhus/.cursor/projects/"
        "remote-edageuclidevhdlbm1-hiran-RTLAssistent-8-CodeAgent-11-AprilBuild-4-case/"
        "assets/"
        "c__Users_madhus_AppData_Roaming_Cursor_User_workspaceStorage_"
        "b1927724be10d68e6005b026951a89e3_images_image-4cffeeec-efdc-44b2-b735-"
        "15c608b7bd5b.png"
    ),
    HERE.parent / "signature.png",
]

OUT = HERE.parent / "ALL_WEEKS_COMBINED.docx"
START_DATE = date(2025, 9, 23)

FIXED = {
    0: "2425488",
    1: "20241504",
    2: "Hiran Abeywardhana",
    3: "MSc Big Data Analytics",
    4: "CMM799 MSc Project",
    5: (
        "Cross-Domain Machine Learning Framework for Scalable GUI "
        "Element Detection and Adaptation in Desktop Environments"
    ),
    6: "Mr.Pumudu Fernando",
}

# Combined doc: slightly more room for progress/plan on a full page
MAX_PROGRESS_WORDS = 72
MAX_PLAN_WORDS = 28
MAX_MEETING_NOTE_WORDS = 24
BODY_FONT_PT = 10
SIG_WIDTH_IN = 0.65

ROW_HEIGHTS_IN = {
    0: 0.24,
    1: 0.24,
    2: 0.24,
    3: 0.24,
    4: 0.24,
    5: 0.45,
    6: 0.28,
    7: 0.28,
    8: 0.24,
    9: 0.55,
    10: 3.4,
    11: 1.6,
    12: 0.28,
    13: 0.55,
}


def find_signature() -> Path:
    for p in SIGNATURE_CANDIDATES:
        if p.is_file():
            return p
    raise FileNotFoundError("signature PNG not found; copy to docs/weekly_reports/signature.png")


def condense(text: str, max_words: int) -> str:
    text = re.sub(r"\s+", " ", text.strip())
    words = text.split()
    if len(words) <= max_words:
        return text
    return " ".join(words[:max_words])


def set_cell_font(cell, size_pt: int) -> None:
    for para in cell.paragraphs:
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        for run in para.runs:
            run.font.size = Pt(size_pt)


def set_cell_text(cell, text: str, size_pt: int = BODY_FONT_PT) -> None:
    cell.text = text
    set_cell_font(cell, size_pt)


def clone_table_into(doc: Document, template: Document) -> Table:
    tbl_el = deepcopy(template.tables[0]._element)
    doc.element.body.append(tbl_el)
    table = Table(tbl_el, doc)
    compact_table(table)
    return table


def compact_table(table: Table) -> None:
    for i, row in enumerate(table.rows):
        h = ROW_HEIGHTS_IN.get(i, 0.22)
        row.height = Inches(h)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


def fill_week_table(table: Table, week_num: int, sig_path: Path) -> None:
    tue = START_DATE + timedelta(weeks=week_num - 1)
    rep = tue + timedelta(weeks=1)
    w = WEEKS[week_num]

    note = w["note"] if w["meeting"] else "NA"
    if w["meeting"]:
        note = condense(note, MAX_MEETING_NOTE_WORDS)

    variable = {
        7: tue.strftime("%d/%m/%Y"),
        8: "Yes" if w["meeting"] else "No",
        9: note,
        10: condense(w["progress"], MAX_PROGRESS_WORDS),
        11: condense(w["plan"], MAX_PLAN_WORDS),
        12: rep.strftime("%d/%m/%Y"),
    }

    for row, val in {**FIXED, **variable}.items():
        set_cell_text(table.cell(row, 1), val, BODY_FONT_PT)

    sig_cell = table.cell(13, 1)
    sig_cell.text = ""
    sig_run = sig_cell.paragraphs[0].add_run()
    sig_run.add_picture(str(sig_path), width=Inches(SIG_WIDTH_IN))


def main() -> None:
    if not TEMPLATE.is_file():
        raise FileNotFoundError(TEMPLATE)
    sig = find_signature()

    template = Document(str(TEMPLATE))
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    for week_num in range(1, 45):
        table = clone_table_into(doc, template)
        if week_num > 1:
            table.rows[0].cells[0].paragraphs[0].paragraph_format.page_break_before = True
        fill_week_table(table, week_num, sig)

    doc.save(str(OUT))
    print(f"OK wrote {OUT}")
    print("Open in Word -> Save As PDF. Expect ~44 pages.")


if __name__ == "__main__":
    main()
