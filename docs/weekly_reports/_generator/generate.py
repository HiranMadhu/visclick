"""Bulk-generate 44 weekly progress reports using template-copy + field-fill.

Usage: python3 generate.py
"""
import shutil
import sys
from datetime import date, timedelta
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
from plan_data import WEEKS

from docx import Document
from docx.shared import Inches

TEMPLATE = "/u/madhus/Week 1 (002).docx"
SIGNATURE = (
    "/u/madhus/.cursor/projects/"
    "remote-edageuclidevhdlbm1-hiran-RTLAssistent-8-CodeAgent-11-AprilBuild-4-case/"
    "assets/"
    "c__Users_madhus_AppData_Roaming_Cursor_User_workspaceStorage_"
    "b1927724be10d68e6005b026951a89e3_images_image-4cffeeec-efdc-44b2-b735-"
    "15c608b7bd5b.png"
)
OUT_DIR = Path(
    "/remote/edageuclidevhdlbm1/hiran/RTLAssistent/8-CodeAgent/11-AprilBuild/"
    "4-case/gui_temp/visclick/docs/weekly_reports"
)
START_DATE = date(2025, 9, 23)  # Tuesday

FIXED = {
    0: "2425488",
    1: "20241504",
    2: "Hiran Abeywardhana",
    3: "MSc Big Data Analytics",
    4: "CMM799 MSc Project",
    5: ("Cross-Domain Machine Learning Framework for Scalable GUI "
        "Element Detection and Adaptation in Desktop Environments"),
    6: "Mr.Pumudu Fernando",
}


def build_one(week_num: int) -> Path:
    tue = START_DATE + timedelta(weeks=week_num - 1)
    rep = tue + timedelta(weeks=1)
    week_commencing = tue.strftime("%d/%m/%Y")
    date_of_report = rep.strftime("%d/%m/%Y")

    w = WEEKS[week_num]
    variable = {
        7: week_commencing,
        8: "Yes" if w["meeting"] else "No",
        9: w["note"],
        10: w["progress"],
        11: w["plan"],
        12: date_of_report,
    }

    out_path = OUT_DIR / f"Week_{week_num:02d}_{tue.strftime('%d_%b_%Y')}.docx"
    shutil.copy(TEMPLATE, out_path)
    doc = Document(out_path)
    t = doc.tables[0]
    for row, val in {**FIXED, **variable}.items():
        t.cell(row, 1).text = val

    # Signature at 0.75 inch wide, inline in row 13
    sig_cell = t.cell(13, 1)
    sig_cell.text = ""
    sig_run = sig_cell.paragraphs[0].add_run()
    sig_run.add_picture(SIGNATURE, width=Inches(0.75))

    doc.save(out_path)
    return out_path


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    generated = []
    for week_num in range(1, 45):
        p = build_one(week_num)
        generated.append(p)
        print(f"  W{week_num:02d}  {p.name}")

    print()
    print(f"OK generated {len(generated)} files in {OUT_DIR}")


if __name__ == "__main__":
    main()
